import pandas as pd
import docx
import csv
from transformers import AutoTokenizer, AutoModelForSequenceClassification
from torch.nn.functional import softmax
import nltk
from nltk.sentiment import SentimentIntensityAnalyzer
from nltk.tokenize import sent_tokenize
from google.cloud import datastore

# Required nltk files, only downloaded on first use
nltk.download('vader_lexicon')
nltk.download('punkt')

# FRONTEND SHOULD ONLY INTERFACE WITH Handler.process_request()
# FILE HANDLING AND OTHER SPECIFICS MAY NEED TO CHANGE DEPENDING ON FRONTEND

# TODO: Add configuration file (for database key, weighting, datastore namespace, sentiment model)
# TODO: SentimentAnalyzer.weight_sentiment() implementation
# TODO: Implement Summary class
# TODO: Implement standard for constants in keyword file (eg headers, like "Keyword") and datastore entries (eg names of properties)


# Handles the process of receiving a transcript, analyzing for sentiment,
# and uploading results to datastore
class Handler:
    def __init__(self):
        self.sentiment_analyzer = SentimentAnalyzer()
        self.database = Database("key.json")
        self.transcript_processor = None
        self.keyword_analyzer = None
        self.data_manager = None

    def process_request(self, transcript_file_path, keywords_file_path, company, date):
        self.transcript_processor = TranscriptProcessor(transcript_file_path)
        self.keyword_analyzer = KeywordAnalyzer(keywords_file_path)
        self.data_manager = DataManager()

        paragraphs = self.transcript_processor.get_paragraphs()
        for paragraph in paragraphs:
            self.process_paragraph(paragraph)

        for data in self.data_manager.get_data():
            self.database.create_entity("Test", data)

        self.data_manager.save_as_csv("test.csv")

    def process_paragraph(self, paragraph):
        found_keywords = self.keyword_analyzer.find_keywords(paragraph)
        
        if not found_keywords:
            return

        score, magnitude = self.sentiment_analyzer.analyze_sentiment(paragraph)
        for entry in found_keywords:
            weight = self.keyword_analyzer.get_weight(entry)
            weighted_score = self.sentiment_analyzer.weight_sentiment(score.item(), weight)

            self.data_manager.add_data(
                Paragraph = paragraph,
                Score = score.item(),
                Magnitude = magnitude,
                WeightedScore = weighted_score,
                **entry
            )


# Opens and parses transcript into paragraphs
# Paragraphs are stored as list of strings
class TranscriptProcessor:
    def __init__(self, file_path):
        self.file_path = file_path
        self.document = docx.Document(file_path)
        self.paragraphs = self.split_paragraphs()

    def split_paragraphs(self):
        '''
            Args: None

            Returns:
                self.document as a list of a paragraphs
        '''
        paragraphs = []
        for para in self.document.paragraphs:
            if para.text:
                paragraphs.append(para.text)
        return paragraphs

    def get_paragraphs(self):
        ''' Getter for self.paragraphs '''
        return self.paragraphs


# Opens and saves keywords as list of dicts
class KeywordAnalyzer:
    def __init__(self, file_path):
        self.file_path = file_path
        self.keywords = self.read_file()
        self.importance_weights = { 'Very Important': 1.5,
                                    'Important': 1.0,
                                    'Less so important': 0.5 } 

    def read_file(self):
        '''
            Returns:
                The keyword excel file at self.file_path saved as a list of dictionaries.
                Example dictionary: {Keyword: 'Assets', Category: 'Business', ...}
        '''
        return pd.read_excel(self.file_path).to_dict('records')

    def find_keywords(self, text):
        '''
            Args:
                text: string of text

            Returns:
                List of dictionaries for all keywords found in text.
        '''
        found = []
        text = text.lower()
        for entry in self.keywords:
            keyword = entry.get('Keyword')
            if keyword.lower() in text:
                found.append(entry)
        return found

    def get_weight(self, entry):
        '''
            Args:
                entry: dictionary entry for a keyword

            Returns:
                The weight for that keyword if there is one, and None otherwise
        '''
        weight = entry.get('Weight')
        if weight in self.importance_weights:
            return self.importance_weights.get(weight)
        else:
            return weight


# Analyzes sentiment of text (does not store any data)
class SentimentAnalyzer:
    def __init__(self):
        self.model = AutoModelForSequenceClassification.from_pretrained('ProsusAI/finBERT')
        self.tokenizer = AutoTokenizer.from_pretrained('ProsusAI/finBERT')
        self.sia = SentimentIntensityAnalyzer()

    def get_probabilities(self, text):
        '''
            Args:
                text: string of text

            Returns:
                Sentiment probabilities of 'text'
                Appears to be a list of [positive, negative, neutral] sentiment values between 0 and 1
        '''
        inputs = self.tokenizer(text, return_tensors="pt", padding=True, truncation=True)
        outputs = self.model(**inputs)
        probabilities = softmax(outputs.logits, dim=1)
        return probabilities[0]

    def analyze_sentiment(self, text):
        '''
            Args:
                text: string of text, presumably a paragraph

            Returns:
                Tuple of (sentiment_score, total_magnitude)
                sentiment_score: sentiment of 'text'
                total_magnitude: magnitude of 'text' sentiment
        '''
        probabilities = self.get_probabilities(text)
        sentiment_score = (probabilities[1] + (probabilities[2] * 2) + (probabilities[0] * 3)) - 2
        
        sentences = sent_tokenize(text)
        magnitudes = []
        for sentence in sentences:
            sentence_magnitude = abs(self.sia.polarity_scores(sentence)['compound'])
            magnitudes.append(sentence_magnitude)
        total_magnitude = sum(magnitudes)
        
        return sentiment_score, total_magnitude

    def weight_sentiment(self, sentiment, weight):
        '''
            Args:
                sentiment:  sentiment score
                weight:     weight to apply to 'sentiment' score 

            Returns:
                Weighted sentiment score (ie 'sentiment' score scaled/adjusted by 'weight')
        '''
        # TODO: implement weighting
        if weight is None:
            return None
        return sentiment * weight


# Stores (sentiment) data as list of dicts, one dict per keyword per paragraph
# Each dict contains at least "Keyword", "Paragraph", "Magnitude", and "Score"
class DataManager:
    def __init__(self):
        self.data = []

    def add_data(self, **kwargs):
        '''
            Args:
                kwargs: dictionary containing a paragraph, the keyword found in that found paragraph,
                the paragraph sentiment score, and related information

            Results:
                Appends dictionary to self.data
        '''
        self.data.append(kwargs)
    
    def get_data(self):
        ''' Getter for self.data '''
        return self.data

    def save_as_csv(self, target_file_path):
        '''
            Args:
                target_file_path: File path for new CSV
            
            Results:
                Saves the data in self.data as a CSV
        '''
        field_names = self.data[0].keys()

        # Open the file in write mode
        with open(target_file_path, mode='w', newline='') as file:
            writer = csv.DictWriter(file, fieldnames=field_names)
            writer.writeheader()
            writer.writerows(self.data)


# Interfaces with datastore
class Database:
    def __init__(self, key_filepath):
        # Client is currently initialized using json file with key
        # If environment variable is set, use: datastore.Client()
        self.client = datastore.Client.from_service_account_json(key_filepath)

    def create_entity(self, kind, data, namespace=None):
        '''
            Args:
                kind:       string containing datastore 'kind' for 'data'
                data:       data to upload to datastore
                namespace:  string containing datastore 'namespace' for 'data'
        
            Results:
                Uploads 'data' to datastore under corresponding 'kind' and 'namespace'
        '''
        entity = self.client.entity(self.client.key(kind, namespace=namespace))
        entity.update(data)
        self.client.put(entity=entity)

    
    
